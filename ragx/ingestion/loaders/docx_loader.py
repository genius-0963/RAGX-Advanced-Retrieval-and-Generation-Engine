"""
RAGX DOCX Loader — Extract text from Word documents using python-docx.

Reads paragraphs and tables, tracking section headings via paragraph
styles.  Each paragraph becomes a :class:`Document` with its section
context preserved in metadata.
"""

from __future__ import annotations

from pathlib import Path
from typing import ClassVar

from ragx.config.logging_config import get_logger
from ragx.ingestion.loaders.base import BaseLoader, Document

logger = get_logger(__name__)


class DocxLoader(BaseLoader):
    """Loader for Microsoft Word ``.docx`` (and ``.doc``) files.

    Paragraphs are grouped by their heading context so that section
    information is available to downstream processors.  Tables are
    serialised row-by-row as pipe-delimited text.
    """

    supported_extensions: ClassVar[list[str]] = [".docx", ".doc"]

    def load(self, source: str) -> list[Document]:
        """Load a DOCX file and return one document per paragraph/table.

        Args:
            source: File-system path to a ``.docx`` file.

        Returns:
            A list of :class:`Document` instances.

        Raises:
            FileNotFoundError: If *source* does not exist.
        """
        from docx import Document as DocxDocument  # type: ignore[import-untyped]

        path = Path(source).resolve()
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        logger.info("Loading DOCX", path=str(path))

        try:
            docx_doc = DocxDocument(str(path))
        except Exception as exc:
            logger.error("Failed to open DOCX", path=str(path), error=str(exc))
            raise

        documents: list[Document] = []
        current_section: str = ""
        paragraph_index: int = 0

        # ── Process paragraphs ──────────────────────────────────────────
        for para in docx_doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""

            # Update current section when a heading is encountered
            if style_name.startswith("heading"):
                current_section = para.text.strip()
                # Also emit the heading itself as a document
                if current_section:
                    documents.append(
                        Document(
                            content=current_section,
                            metadata={
                                "paragraph_index": paragraph_index,
                                "section": current_section,
                                "style": style_name,
                                "element_type": "heading",
                                "file_name": path.name,
                            },
                            source_path=str(path),
                        )
                    )
                    paragraph_index += 1
                continue

            text = para.text.strip()
            if not text:
                continue

            documents.append(
                Document(
                    content=text,
                    metadata={
                        "paragraph_index": paragraph_index,
                        "section": current_section,
                        "style": style_name,
                        "element_type": "paragraph",
                        "file_name": path.name,
                    },
                    source_path=str(path),
                )
            )
            paragraph_index += 1

        # ── Process tables ──────────────────────────────────────────────
        for table_idx, table in enumerate(docx_doc.tables):
            rows_text: list[str] = []
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cell_texts))

            table_content = "\n".join(rows_text)
            if not table_content.strip():
                continue

            documents.append(
                Document(
                    content=table_content,
                    metadata={
                        "paragraph_index": paragraph_index,
                        "section": current_section,
                        "element_type": "table",
                        "table_index": table_idx,
                        "file_name": path.name,
                    },
                    source_path=str(path),
                )
            )
            paragraph_index += 1

        logger.info(
            "DOCX loaded",
            path=str(path),
            documents=len(documents),
        )
        return documents
